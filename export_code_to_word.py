import os
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

PROJECT_ROOT = os.getcwd()
OUTPUT_FILE = "Resume_Screening_System_FULL_SOURCE_CODE.docx"

ALLOWED_EXTENSIONS = (
    ".py", ".ts", ".tsx", ".js", ".json",
    ".prisma", ".md", ".mjs", ".cjs"
)

EXCLUDE_DIRS = {
    "node_modules", ".next", ".git", "__pycache__", "venv"
}

doc = Document()
doc.add_heading("Resume Screening System – Full Source Code", level=1)

for root, dirs, files in os.walk(PROJECT_ROOT):
    dirs[:] = [d for d in dirs if d not in EXCLUDE_DIRS]

    for file in files:
        if file.endswith(ALLOWED_EXTENSIONS):
            file_path = os.path.join(root, file)
            relative_path = os.path.relpath(file_path, PROJECT_ROOT)

            # Add file heading
            doc.add_heading(relative_path, level=2)

            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    code = f.read()
            except Exception as e:
                code = f"Could not read file: {e}"

            # Add code paragraph
            para = doc.add_paragraph()
            run = para.add_run(code)
            run.font.name = "Courier New"
            run.font.size = Pt(10)

            # Force monospace in Word
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')

# Save Word file
doc.save(OUTPUT_FILE)
print(f"✅ Word file created: {OUTPUT_FILE}")
